from pathlib import Path

from docx import Document
from docx.shared import Pt


DOCX = Path("TCC_OdontoPlay.docx")


def set_para(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def set_cell(cell, text):
    cell.text = text


def resize_table_rows(table, row_count):
    while len(table.rows) < row_count:
        table.add_row()
    while len(table.rows) > row_count:
        tr = table.rows[-1]._tr
        tr.getparent().remove(tr)


def remove_extra_columns(table, col_count):
    while len(table.columns) > col_count:
        idx = len(table.columns) - 1
        for row in table.rows:
            tc = row._tr.tc_lst[idx]
            row._tr.remove(tc)
        tbl_grid = table._tbl.tblGrid
        if tbl_grid is not None and len(tbl_grid.gridCol_lst) > idx:
            tbl_grid.remove(tbl_grid.gridCol_lst[idx])


def fill_table(table, data):
    resize_table_rows(table, len(data))
    remove_extra_columns(table, len(data[0]))
    for r, row in enumerate(data):
        for c, value in enumerate(row):
            set_cell(table.cell(r, c), value)


def set_table_font(table, size_pt):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(size_pt)


doc = Document(DOCX)

functional_requirements = [
    [
        "RF001",
        "Cadastro da Criança",
        "O aplicativo deve permitir o cadastro de crianças que o utilizarão.\n\n"
        "Campos do cadastro:\n"
        "1. Nome da criança: String\n"
        "2. Idade: Integer\n"
        "3. Data de nascimento: Date\n"
        "4. Sexo: String\n"
        "5. Nome do responsável: String\n"
        "6. Contato do responsável: String\n"
        "7. Data de cadastro: Date\n\n"
        "Prioridade: Essencial",
    ],
    [
        "RF002",
        "Login",
        "O aplicativo deve permitir que a criança ou responsável acesse o aplicativo utilizando credenciais previamente cadastradas, possibilitando o acesso às funcionalidades do sistema.\n\n"
        "Campos do login:\n"
        "1. E-mail: String\n"
        "   - Campo obrigatório\n"
        "   - Deve seguir formato válido de e-mail, por exemplo: usuario@email.com\n"
        "2. Senha: String\n"
        "   - Campo obrigatório\n"
        "   - Deve possuir no mínimo 6 caracteres\n\n"
        "Prioridade: Essencial",
    ],
    [
        "RF003",
        "Jogo I - Consultório Odontológico Virtual",
        "A criança poderá explorar um consultório odontológico virtual onde um personagem dentista apresentará os instrumentos de forma amigável. A criança poderá clicar nos instrumentos, ver imagens e ouvir explicações.\n\n"
        "Exemplos:\n"
        "Sugador: \"Esse canudinho suga a água da boca.\"\n"
        "Espelho odontológico: \"Esse espelhinho ajuda o dentista a ver os dentes.\"\n\n"
        "Prioridade: Essencial",
    ],
    [
        "RF004",
        "Jogo II",
        "Jogo II - Jogo de Higiene Bucal.\n\n"
        "A criança assume o papel de dentista e deve limpar os dentes de um personagem, aprendendo os passos corretos da higiene bucal.\n\n"
        "Etapas:\n"
        "1. Identificar dentes sujos\n"
        "2. Usar escova dental\n"
        "3. Aplicar pasta de dente\n"
        "4. Enxaguar os dentes\n\n"
        "Prioridade: Essencial",
    ],
    [
        "RF005",
        "Sistema de recompensa",
        "O aplicativo deve possuir um sistema de recompensas para incentivar o engajamento da criança durante o uso dos jogos.\n\n"
        "Elemento de recompensa:\n"
        "- Estrelas: as estrelas devem ser concedidas quando a criança completar jogos ou atividades dentro do aplicativo.\n\n"
        "Prioridade: Desejável",
    ],
    [
        "RF006",
        "Acesso aos jogos",
        "O aplicativo deve permitir que a criança visualize, selecione e acesse os jogos disponíveis por meio da tela principal.\n\n"
        "Funcionalidades:\n"
        "1. Exibir lista de jogos disponíveis\n"
        "2. Permitir selecionar o jogo desejado\n"
        "3. Iniciar o jogo selecionado\n\n"
        "Prioridade: Essencial",
    ],
    [
        "RF007",
        "Áudio narrado",
        "O aplicativo deve permitir que os jogos possuam narração em áudio explicando os instrumentos e as atividades.\n\n"
        "Funcionalidades:\n"
        "1. Reproduzir áudio explicativo ao clicar em instrumentos\n"
        "2. Reproduzir instruções durante os jogos\n\n"
        "Prioridade: Desejável",
    ],
    [
        "RF008",
        "Escolha de personagem",
        "O sistema deve disponibilizar uma tela de seleção de personagem, permitindo que o usuário escolha entre um personagem masculino ou feminino antes de iniciar os Jogos 1 e 2. A escolha realizada deverá ser aplicada durante a execução do respectivo jogo.\n\n"
        "Prioridade: Importante",
    ],
    [
        "RF009",
        "Controle de Música Tema",
        "O sistema deve disponibilizar, na tela principal de seleção de jogos, um botão que permita ao usuário ativar ou desativar a música tema do aplicativo. A configuração selecionada deverá ser aplicada imediatamente e permanecer válida durante a navegação entre as telas do sistema.\n\n"
        "Prioridade: Importante",
    ],
    [
        "RF010",
        "Gerenciamento Administrativo",
        "O sistema deve permitir que administradores visualizem e gerenciem os dados cadastrados no aplicativo por meio da plataforma web.\n\n"
        "Funcionalidades:\n"
        "1. Visualizar crianças cadastradas\n"
        "2. Consultar progresso e pontuação\n"
        "3. Editar ou remover registros quando necessário\n\n"
        "Prioridade: Importante",
    ],
]

fill_table(doc.tables[0], [["ID", "Requisito", "Descrição"]] + functional_requirements)
set_table_font(doc.tables[0], 8.5)

fill_table(
    doc.tables[2],
    [
        ["ID", "Perigo", "Soluções"],
        [
            "RS001",
            "No RF001 e RF010, dados pessoais da criança e do responsável podem ser expostos, alterados indevidamente ou removidos sem autorização.",
            "Aplicar autenticação, regras de permissão no Firebase, validação dos campos e controle de acesso administrativo na plataforma web.",
        ],
        [
            "RS002",
            "No RF002, credenciais fracas ou inválidas podem impedir o acesso correto ou permitir tentativa de uso indevido da conta.",
            "Exigir e-mail em formato válido, senha com no mínimo 6 caracteres, mensagens de erro controladas e autenticação gerenciada pelo Firebase.",
        ],
        [
            "RS003",
            "No RF003, RF004, RF007 e RF009, conteúdo visual ou sonoro pode gerar desconforto ou ansiedade em crianças sensíveis.",
            "Usar linguagem acolhedora, narração simples, representação amigável dos instrumentos, instruções adequadas à idade e opção de controle da música tema.",
        ],
        [
            "RS004",
            "No RF005 e RF010, falhas de gravação podem comprometer recompensas, progresso, pontuação e consultas administrativas.",
            "Registrar progresso e pontuação de forma persistente, validar gravações e manter consistência entre estrelas, jogos concluídos e relatórios administrativos.",
        ],
        [
            "RS005",
            "No RF006 e RF008, a criança pode acessar jogos fora do fluxo previsto ou iniciar atividades sem selecionar o personagem adequado.",
            "Organizar a seleção de jogos em interface controlada, validar disponibilidade dos assets e aplicar a escolha de personagem antes da execução dos Jogos 1 e 2.",
        ],
    ],
)
set_table_font(doc.tables[2], 8.5)

fill_table(
    doc.tables[3],
    [
        ["Segurança/\nprobabilidade", "Catastrófico\n(1)", "Crítico\n(2)", "Marginal\n(3)", "Insignificante\n(4)"],
        ["Frequente\n(A)", "", "", "RS003", ""],
        ["Provável\n(B)", "", "RS001", "", ""],
        ["Ocasional\n(C)", "", "RS004", "RS002", ""],
        ["Remoto\n(D)", "", "", "RS005", ""],
        ["Improvável\n(E)", "", "", "", ""],
        ["Eliminado\n(F)", "", "", "", ""],
    ],
)
set_table_font(doc.tables[3], 8)

trace_header = ["ID de requisitos"] + [f"RF{i:03d}" for i in range(1, 11)]
relations = {
    "RF001": {"RF002": "D", "RF010": "D"},
    "RF002": {"RF001": "D", "RF006": "D", "RF010": "D"},
    "RF003": {"RF006": "D", "RF007": "D", "RF008": "D", "RF005": "D"},
    "RF004": {"RF006": "D", "RF007": "D", "RF008": "D", "RF005": "D"},
    "RF005": {"RF003": "D", "RF004": "D", "RF010": "D"},
    "RF006": {"RF002": "D", "RF003": "D", "RF004": "D", "RF008": "D", "RF009": "D"},
    "RF007": {"RF003": "D", "RF004": "D", "RF009": "D"},
    "RF008": {"RF003": "D", "RF004": "D", "RF006": "D"},
    "RF009": {"RF006": "D", "RF007": "D"},
    "RF010": {"RF001": "D", "RF002": "D", "RF005": "D"},
}
trace_rows = [trace_header]
for i in range(1, 11):
    rf = f"RF{i:03d}"
    trace_rows.append([rf] + [relations.get(rf, {}).get(f"RF{j:03d}", "") for j in range(1, 11)])
fill_table(doc.tables[4], trace_rows)
set_table_font(doc.tables[4], 7.5)

fill_table(
    doc.tables[5],
    [
        ["REQUISITO FUNCIONAL", "COMPLEXIDADE", "ALI", "AIE", "SE", "EE", "CE", "TOTAL"],
        ["RF001", "MÉDIA", "10", "0", "0", "4", "3", "17"],
        ["RF002", "MÉDIA", "0", "5", "0", "4", "3", "12"],
        ["RF003", "MÉDIA", "0", "0", "5", "4", "3", "12"],
        ["RF004", "MÉDIA", "0", "0", "5", "4", "3", "12"],
        ["RF005", "BAIXA", "7", "0", "4", "3", "0", "14"],
        ["RF006", "BAIXA", "0", "0", "4", "3", "3", "10"],
        ["RF007", "BAIXA", "0", "0", "4", "3", "0", "7"],
        ["RF008", "BAIXA", "0", "0", "4", "3", "0", "7"],
        ["RF009", "BAIXA", "0", "0", "4", "3", "0", "7"],
        ["RF010", "ALTA", "15", "5", "7", "6", "6", "39"],
        ["TOTAL ESTIMADO", "-", "32", "10", "33", "37", "21", "137"],
    ],
)
set_table_font(doc.tables[5], 8)

replacements = {
    "cadastro da criança, edição/exclusão de cadastro, login, acesso aos jogos, jogos educativos, áudio narrado, recompensas, registro e visualização de progresso": "cadastro da criança, login, jogos educativos, sistema de recompensas, acesso aos jogos, áudio narrado, escolha de personagem, controle de música tema e gerenciamento administrativo",
    "cadastro da criança, login, acesso aos jogos, jogos educativos, áudio narrado, recompensas, registro e visualização de progresso": "cadastro da criança, login, jogos educativos, sistema de recompensas, acesso aos jogos, áudio narrado, escolha de personagem, controle de música tema e gerenciamento administrativo",
    "cadastro da criança, edição ou exclusão de cadastro, login, seleção de jogos, interação com instrumentos, execução da higiene bucal, áudio narrado e conclusão de atividades": "cadastro da criança, login, seleção de jogos, interação com instrumentos, execução da higiene bucal, escolha de personagem, áudio narrado, controle de música tema e gestão administrativa dos dados",
    "explicações narradas, mensagens de conclusão, estrelas, feedbacks educativos e estados visuais decorrentes da interação da criança com o consultório virtual e o jogo de higiene bucal": "imagens, explicações narradas, mensagens de conclusão, estrelas, feedbacks educativos, música tema e estados visuais decorrentes da interação da criança com o consultório virtual e o jogo de higiene bucal",
    "carregamento de cadastro, exibição da lista de jogos, visualização de progresso, jogos concluídos, estrelas obtidas e histórico de atividades": "carregamento de cadastro, exibição da lista de jogos, seleção de personagem, estado da música tema, progresso, pontuação e dados administrativos consultados na plataforma web",
    "cadastro da criança, edição/exclusão de cadastro, login, seleção de jogos, consultório odontológico virtual, jogo de higiene bucal, áudio narrado, recompensas, registro e visualização de progresso": "cadastro da criança, login, seleção de jogos, consultório odontológico virtual, jogo de higiene bucal, sistema de recompensa, áudio narrado, escolha de personagem, controle de música tema e gerenciamento administrativo via plataforma web",
    "cadastro da criança, login, seleção de jogos, consultório odontológico virtual, jogo de higiene bucal, áudio narrado, recompensas, registro e visualização de progresso": "cadastro da criança, login, seleção de jogos, consultório odontológico virtual, jogo de higiene bucal, sistema de recompensa, áudio narrado, escolha de personagem, controle de música tema e gerenciamento administrativo via plataforma web",
    "autenticação, cadastro da criança, progresso e recompensas": "autenticação, cadastro da criança, recompensas, pontuação e dados administrativos",
    "registro e visualização de progresso": "controle de recompensas, pontuação e gerenciamento administrativo",
    "controle de música": "controle de música tema",
}

for paragraph in doc.paragraphs:
    new_text = paragraph.text
    for old, new in replacements.items():
        new_text = new_text.replace(old, new)
    if new_text != paragraph.text:
        set_para(paragraph, new_text)

doc.save(DOCX)
print(DOCX)
