from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph


DOCX = Path("TCC_OdontoPlay2.docx")


def set_para(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = paragraph._p.__class__()
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_para.style = style
    new_para.add_run(text)
    return new_para


doc = Document(DOCX)

target = None
for i, paragraph in enumerate(doc.paragraphs):
    if paragraph.text.strip() == "4.1 Pré-Requisitos":
        target = doc.paragraphs[i + 1]
        break

if target is None:
    raise RuntimeError("Seção 4.1 Pré-Requisitos não encontrada.")

set_para(
    target,
    "Para utilização do OdontoPlay pelo público final, é necessário dispor de um smartphone ou tablet compatível com aplicativos móveis Android ou iOS, acesso à internet para autenticação e sincronização dos dados, conta previamente cadastrada no sistema e permissão para reprodução de áudio, uma vez que parte da experiência educativa utiliza narrações, efeitos sonoros e música tema.",
)

p = insert_paragraph_after(
    target,
    "No contexto de uso, não é necessário que a criança ou o responsável instale ferramentas de programação. A navegação ocorre diretamente pela interface do aplicativo, contemplando cadastro, login, escolha de personagem, seleção de jogos, execução das atividades educativas, controle da música tema e visualização dos recursos de recompensa.",
    target.style,
)
p = insert_paragraph_after(
    p,
    "Para desenvolvimento, manutenção ou continuidade acadêmica do projeto, recomenda-se um computador com ambiente React Native/Expo configurado, Node.js em versão compatível com o Expo, gerenciador npm, Expo CLI ou execução via npx, dispositivo físico com Expo Go ou emulador Android/iOS, acesso à internet e projeto Firebase devidamente configurado. A aplicação utiliza React Native, Expo, TypeScript, React Navigation, Firebase Authentication e Firestore.",
    target.style,
)

doc.save(DOCX)
print(DOCX)
