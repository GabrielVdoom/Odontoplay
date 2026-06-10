from pathlib import Path

from docx import Document
from docx.shared import Pt


DOCX = Path("TCC_OdontoPlay_registro_software.docx")
OUT = Path("TCC_OdontoPlay_registro_software_adaptado.docx")


def set_para(paragraph, text):
    """Replace paragraph text while preserving the paragraph style."""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def clear_para(paragraph):
    set_para(paragraph, "")


def set_cell(cell, text):
    cell.text = text


def resize_table_rows(table, row_count):
    while len(table.rows) < row_count:
        table.add_row()
    while len(table.rows) > row_count:
        tr = table.rows[-1]._tr
        tr.getparent().remove(tr)


def fill_table(table, data):
    resize_table_rows(table, len(data))
    for r, row in enumerate(data):
        for c, value in enumerate(row):
            set_cell(table.cell(r, c), value)


def set_table_font(table, size_pt):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(size_pt)


doc = Document(DOCX)
p = doc.paragraphs

title = "ODONTOPLAY: APLICATIVO MOBILE GAMIFICADO PARA EDUCAÇÃO ODONTOLÓGICA INFANTIL"

set_para(p[0], "[NOME DA INSTITUIÇÃO]")
set_para(p[47], "[CIDADE-ESTADO]")
set_para(p[95], "[CIDADE-ESTADO]")

for idx in (20, 67, 112):
    set_para(p[idx], title)

for idx in (75, 120):
    set_para(
        p[idx],
        "Trabalho de Conclusão de Curso apresentado como requisito final para conclusão do curso de Sistemas de Informação do [NOME DA INSTITUIÇÃO], sob a orientação do(a) professor(a) [NOME DO(A) ORIENTADOR(A)] (e coorientação do(a) professor(a) [NOME DO(A) COORIENTADOR(A)], se houver).",
    )

set_para(p[129], "Orientador(a): [TITULAÇÃO E NOME DO(A) ORIENTADOR(A)]")
set_para(p[134], "Avaliador(a) interno(a): [TITULAÇÃO E NOME DO(A) AVALIADOR(A)]")
set_para(p[139], "Avaliador(a) interno(a): [TITULAÇÃO E NOME DO(A) AVALIADOR(A)]")

set_para(p[145], "[Texto de agradecimento do Autor 1, a ser preenchido posteriormente.]")
set_para(p[148], "[NOME DO AUTOR 1]")
set_para(p[154], "[Texto de agradecimento do Autor 2, a ser preenchido posteriormente.]")
set_para(p[157], "[NOME DO AUTOR 2]")
set_para(p[203], '"[Frase de efeito opcional relacionada ao projeto ou à trajetória acadêmica.]"')
set_para(p[204], "[Autor da frase]")

set_para(p[207], "Tabela 1 - Requisitos funcionais do OdontoPlay\t[atualizar paginação]")
set_para(p[241], "Quadro 1 - Requisitos não funcionais do OdontoPlay\t[atualizar paginação]")
set_para(p[253], "Figura 1 - Tela de seleção de jogos do OdontoPlay\t[atualizar paginação]")
set_para(p[254], "Figura 2 - Exemplo de atividade gamificada de higiene bucal\t[atualizar paginação]")

set_para(
    p[264],
    "O OdontoPlay é um aplicativo mobile desenvolvido em React Native, com apoio da plataforma Firebase, voltado para crianças de 4 a 8 anos. O sistema propõe uma experiência lúdica de educação odontológica por meio de jogos, personagens, recursos visuais, trilha sonora e atividades interativas que auxiliam a criança a reconhecer instrumentos odontológicos, compreender práticas de higiene bucal e familiarizar-se com o ambiente do consultório.",
)
set_para(
    p[265],
    "O projeto parte do problema da ansiedade infantil associada ao atendimento odontológico e da dificuldade de comunicação de conceitos de saúde bucal para crianças em fase inicial de desenvolvimento cognitivo. Ao transformar conteúdos preventivos em missões, recompensas e interações guiadas, o OdontoPlay busca apoiar famílias e profissionais de odontologia na construção de uma experiência mais acolhedora antes e durante o contato clínico.",
)

set_para(
    p[289],
    "A análise por pontos de função foi considerada para estimar o tamanho funcional do OdontoPlay a partir das funcionalidades percebidas pelo usuário final, especialmente autenticação, cadastro, seleção de jogos, execução de atividades educativas, registro de progresso e consulta de resultados.",
)
set_para(
    p[290],
    "Para fins acadêmicos, a contagem apresentada possui caráter estimativo e utiliza os componentes tradicionais de dados e transações para representar o escopo funcional do aplicativo mobile e dos serviços de apoio mantidos no Firebase.",
)
set_para(p[292], "1.7 Componentes Fundamentais")
set_para(
    p[293],
    "Funções do tipo Dado: correspondem aos grupos de informações persistidas ou consultadas pelo OdontoPlay, como perfis de usuários, progresso das atividades, configurações e dados de autenticação.",
)
set_para(
    p[294],
    "Funções do tipo Transação: correspondem às operações realizadas pelos usuários no aplicativo, incluindo cadastro, login, seleção de jogos, execução de interações, conclusão de fases e recuperação de dados de progresso.",
)
set_para(p[296], "Arquivo Lógico Interno (ALI)")
set_para(
    p[297],
    "No OdontoPlay, os ALIs abrangem dados mantidos dentro da fronteira da aplicação, como usuários cadastrados, preferências, histórico de progresso e registros de desempenho associados às atividades gamificadas.",
)
set_para(p[298], "Arquivo de Interface Externa (AIE)")
set_para(
    p[299],
    "Os AIEs correspondem a serviços externos utilizados pelo aplicativo, especialmente os recursos gerenciados pela plataforma Firebase, como autenticação e infraestrutura de banco de dados em nuvem.",
)
set_para(p[300], "Entrada Externa (EE)")
set_para(
    p[301],
    "As entradas externas compreendem dados inseridos ou acionados pelo usuário, como cadastro, autenticação, escolha de personagem ou jogo, início de fase, toque em elementos interativos e conclusão de atividades.",
)
set_para(p[302], "Saída Externa (SE)")
set_para(
    p[303],
    "As saídas externas ocorrem quando o sistema apresenta resultados processados, mensagens de conclusão, pontuações, estrelas, feedbacks educativos e estados visuais decorrentes da interação da criança com os jogos.",
)
set_para(p[304], "Consulta Externa (CE)")
set_para(
    p[305],
    "As consultas externas envolvem a recuperação de informações sem alteração direta dos dados, como carregamento de progresso, exibição de telas, seleção de atividades disponíveis e consulta de configurações do usuário.",
)
set_para(p[312], "Determinar o tipo de contagem")
set_para(
    p[313],
    "Para o OdontoPlay, utiliza-se a contagem de projeto em desenvolvimento, considerando as funcionalidades disponíveis na primeira versão funcional do aplicativo e seus módulos essenciais de interação infantil, autenticação e persistência de progresso.",
)
set_para(p[315], "Determinar Escopo e Fronteira")
set_para(
    p[316],
    "O escopo contempla o aplicativo mobile, suas telas de acesso, cadastro, seleção de jogos, atividades educativas, consultório virtual, controle de música e integração com Firebase. A fronteira da aplicação limita-se aos módulos implementados no OdontoPlay e aos serviços externos necessários para autenticação e persistência.",
)
set_para(p[318], "Calcular pontos de função não ajustados")
set_para(
    p[319],
    "A estimativa considera os componentes ALI, AIE, EE, SE e CE relacionados aos requisitos funcionais do sistema, classificando-os por complexidade para dimensionar o esforço funcional inicial.",
)
set_para(p[320], "Será utilizada a tabela abaixo para cálculo dos pontos de função do OdontoPlay.")
set_para(
    p[321],
    "Os valores estimados serão utilizados para calcular os pontos de função não ajustados, multiplicando a quantidade de componentes pelo peso de complexidade correspondente.",
)
set_para(p[324], "PONTO DE FUNÇÃO - ODONTOPLAY")

set_para(
    p[329],
    "Esta seção reúne os dados dos autores responsáveis pelo desenvolvimento e documentação do OdontoPlay. Os campos pessoais devem ser preenchidos posteriormente, mantendo a ordem de autoria definida pela equipe.",
)
set_para(
    p[330],
    "Quando aplicável, indicar também orientador, coorientador e demais colaboradores acadêmicos ou técnicos envolvidos na concepção, implementação, validação e documentação do software.",
)
set_para(p[333], "INFORMAÇÕES DO PRODUTO")
set_para(p[337], "Nome do Produto: OdontoPlay")
set_para(p[338], "Linguagem de Programação: TypeScript e JavaScript")
set_para(p[339], "Banco de dados: Firebase Firestore, com suporte de Firebase Authentication")
set_para(
    p[340],
    "Resumo da Aplicabilidade do Software: aplicativo mobile educativo destinado a crianças de 4 a 8 anos, com jogos e recursos multimídia voltados ao aprendizado de higiene bucal, familiarização com o consultório odontológico e redução da ansiedade infantil.",
)
set_para(
    p[341],
    "Qual problema o software pretende resolver: dificuldade de engajamento infantil em conteúdos preventivos de saúde bucal e ansiedade associada à consulta odontológica, especialmente em crianças pequenas que ainda não compreendem procedimentos, instrumentos e rotinas clínicas.",
)
set_para(p[342], "Framework: React Native com Expo")
set_para(
    p[343],
    "Campo de Atuação do software: Educação, saúde digital, odontologia preventiva e apoio ao atendimento odontopediátrico.",
)
set_para(
    p[344],
    "Tipo de Programa: Aplicativo mobile educacional com recursos de gamificação, multimídia, autenticação e persistência de dados.",
)
set_para(p[345], "Código do Registro: [A preencher após submissão/registro, se houver]")
set_para(p[346], "HASH em SHA512: [A preencher após geração do resumo digital do código-fonte]")
set_para(p[347], "Repositório: [Inserir link do repositório do código-fonte]")

set_para(
    p[350],
    "Este manual apresenta os recursos necessários para instalação, configuração e utilização do OdontoPlay, incluindo requisitos de ambiente, dependências principais, execução do aplicativo e descrição dos módulos disponíveis para crianças, responsáveis e equipe de desenvolvimento.",
)
set_para(
    p[354],
    "Para execução em ambiente de desenvolvimento, recomenda-se computador com Node.js compatível com Expo, gerenciador npm, Expo CLI, dispositivo físico com Expo Go ou emulador Android/iOS, acesso à internet e projeto Firebase configurado. O aplicativo utiliza React Native, Expo, TypeScript, React Navigation, Firebase Authentication e Firestore.",
)
set_para(
    p[357],
    "A instalação consiste em clonar ou obter o repositório do projeto, instalar as dependências com npm install, configurar as credenciais do Firebase no arquivo de serviços correspondente e iniciar o ambiente com npm run dev ou npx expo start. Em seguida, o aplicativo pode ser aberto em um dispositivo móvel por QR Code ou em emulador compatível.",
)
set_para(
    p[360],
    "O fluxo principal inicia na tela de splash, seguida de autenticação ou cadastro. Após o acesso, a criança é direcionada à seleção de jogos, onde pode iniciar atividades educativas. O módulo de consultório apresenta elementos relacionados ao ambiente odontológico, enquanto os jogos trabalham reconhecimento de instrumentos, práticas de escovação, interação com personagens, feedback visual, estrelas e conclusão de fases.",
)
manual_extra = {
    361: "Atores do sistema: criança usuária, responsável ou mediador do uso, profissional de odontologia em contexto educativo e administrador/desenvolvedor responsável pela manutenção técnica.",
    362: "Módulo de autenticação: permite cadastro e login de usuários, possibilitando identificação do progresso individual e acesso às atividades.",
    363: "Módulo de seleção de jogos: organiza as atividades disponíveis em cartões visuais adequados à faixa etária, com navegação simples e elementos gráficos infantis.",
    364: "Módulo de educação odontológica: apresenta conteúdos sobre escovação, higiene bucal, cuidado com os dentes e reconhecimento de instrumentos de forma lúdica e não ameaçadora.",
    365: "Módulo de gamificação: utiliza recompensas visuais, estrelas, feedbacks positivos e repetição de fases para estimular engajamento, autonomia e aprendizagem gradual.",
    366: "Módulo de áudio e ambientação: oferece trilha sonora e controle de música para tornar a experiência mais acolhedora e compatível com o público infantil.",
    367: "Persistência de dados: registra informações essenciais do usuário e do progresso em Firebase, respeitando o princípio de minimização de dados para público infantil.",
}
for idx, text in manual_extra.items():
    set_para(p[idx], text)

refs = {
    381: "EXPO. Expo Documentation. Disponível em: https://docs.expo.dev/. Acesso em: 06 jun. 2026.",
    383: "FIREBASE. Firebase Documentation. Disponível em: https://firebase.google.com/docs. Acesso em: 06 jun. 2026.",
    385: "REACT NATIVE. React Native Documentation. Disponível em: https://reactnative.dev/docs/getting-started. Acesso em: 06 jun. 2026.",
    387: "SOMMERVILLE, Ian. Engenharia de Software. 10. ed. São Paulo: Pearson, 2019.",
    389: "PRESSMAN, Roger S.; MAXIM, Bruce R. Engenharia de Software: uma abordagem profissional. 8. ed. Porto Alegre: AMGH, 2016.",
}
for idx, text in refs.items():
    set_para(p[idx], text)

set_para(p[479], "APÊNDICE A - TRECHO REPRESENTATIVO DA ESTRUTURA DO APLICATIVO ODONTOPLAY")
code_lines = [
    'import React from "react";',
    'import { NavigationContainer } from "@react-navigation/native";',
    'import { createNativeStackNavigator } from "@react-navigation/native-stack";',
    "",
    'import Login from "./screens/Login";',
    'import Cadastro from "./screens/Cadastro";',
    'import SelecaoJogos from "./screens/SelecaoJogos";',
    'import Consultorio from "./screens/Consultorio";',
    'import Jogo2 from "./screens/Jogo2";',
    'import Splash from "./screens/Splash";',
    'import { ThemeMusicProvider } from "./contexts/ThemeMusicContext";',
    "",
    "export type RootStackParamList = {",
    "  Splash: undefined;",
    "  Login: undefined;",
    "  Cadastro: undefined;",
    "  SelecaoJogos: undefined;",
    "  Consultorio: undefined;",
    "  Jogo2: undefined;",
    "};",
    "",
    "const Stack = createNativeStackNavigator<RootStackParamList>();",
    "",
    "export default function App() {",
    "  return (",
    "    <ThemeMusicProvider>",
    "      <NavigationContainer>",
    "        <Stack.Navigator screenOptions={{ headerShown: false }}>",
    '          <Stack.Screen name="Splash" component={Splash} />',
    '          <Stack.Screen name="Login" component={Login} />',
    '          <Stack.Screen name="Cadastro" component={Cadastro} />',
    '          <Stack.Screen name="SelecaoJogos" component={SelecaoJogos} />',
    '          <Stack.Screen name="Consultorio" component={Consultorio} />',
    '          <Stack.Screen name="Jogo2" component={Jogo2} />',
    "        </Stack.Navigator>",
    "      </NavigationContainer>",
    "    </ThemeMusicProvider>",
    "  );",
    "}",
]
for i in range(480, 527):
    text = code_lines[i - 480] if i - 480 < len(code_lines) else ""
    set_para(p[i], text)

fill_table(
    doc.tables[0],
    [
        ["ID", "Requisito", "Descrição"],
        ["RF001", "Cadastro de usuário", "Permitir que responsável ou mediador cadastre usuário com dados mínimos necessários, criando identificação para acesso ao aplicativo e acompanhamento de progresso."],
        ["RF002", "Autenticação", "Permitir login seguro por credenciais, integrando Firebase Authentication e encaminhando o usuário autenticado para a experiência principal."],
        ["RF003", "Seleção de jogos", "Disponibilizar tela de seleção com cartões visuais para acesso às atividades educativas, respeitando a linguagem gráfica adequada à faixa etária de 4 a 8 anos."],
        ["RF004", "Consultório virtual", "Apresentar ambiente odontológico ilustrado para familiarizar a criança com instrumentos, personagens e situações comuns da consulta odontológica."],
        ["RF005", "Jogo de reconhecimento", "Permitir interação com elementos relacionados à odontologia, associando instrumentos, ações e práticas corretas por meio de feedback visual."],
        ["RF006", "Jogo de escovação", "Simular etapas de higiene bucal, incluindo aplicação de pasta, escovação, limpeza e conclusão com reforço positivo."],
        ["RF007", "Sistema de recompensas", "Exibir estrelas, mensagens de sucesso e estados de conclusão para estimular permanência, repetição e aprendizagem progressiva."],
        ["RF008", "Controle de áudio", "Permitir ativar ou desativar música de fundo, mantendo ambientação lúdica sem prejudicar a concentração ou o conforto sensorial."],
        ["RF009", "Persistência de progresso", "Registrar no Firebase informações essenciais de avanço e desempenho para continuidade do uso em sessões futuras."],
        ["RF010", "Navegação entre telas", "Permitir transição fluida entre splash, login, cadastro, seleção de jogos, consultório e atividades, sem exibir cabeçalhos técnicos ao usuário infantil."],
        ["RF011", "Feedback educativo", "Apresentar mensagens, imagens e estados visuais que reforcem boas práticas de higiene bucal e reduzam interpretações ameaçadoras do contexto odontológico."],
        ["RF012", "Validação de formulários", "Validar campos de cadastro e login, informando erros de modo claro para o responsável ou mediador."],
        ["RF013", "Pré-carregamento de assets", "Carregar imagens e recursos gráficos essenciais para reduzir atrasos perceptíveis durante as atividades infantis."],
    ],
)
set_table_font(doc.tables[0], 9)

fill_table(
    doc.tables[1],
    [
        ["ID", "Requisito", "Descrição", "Prioridade"],
        ["RNF01", "Usabilidade infantil", "A interface deve utilizar elementos visuais claros, botões grandes, linguagem simples, contraste adequado e fluxo de baixa complexidade para crianças de 4 a 8 anos.", "Essencial"],
        ["RNF02", "Desempenho", "As telas e assets devem carregar com fluidez suficiente para evitar frustração, interrupção da atividade ou perda de engajamento infantil.", "Essencial"],
        ["RNF03", "Compatibilidade mobile", "O sistema deve funcionar em dispositivos móveis compatíveis com React Native/Expo, considerando variações de resolução e orientação.", "Importante"],
        ["RNF04", "Manutenibilidade", "O código deve ser organizado em telas, serviços e contextos, favorecendo manutenção evolutiva, inclusão de novos jogos e ajustes de conteúdo.", "Importante"],
        ["RNF05", "Acessibilidade e conforto", "O aplicativo deve evitar estímulos excessivos, permitir controle de áudio e utilizar feedbacks positivos, respeitando diferentes níveis de sensibilidade infantil.", "Importante"],
    ],
)
set_table_font(doc.tables[1], 9)

fill_table(
    doc.tables[2],
    [
        ["ID", "Perigo", "Soluções"],
        ["RS001", "Exposição indevida de dados de crianças ou responsáveis.", "Coletar apenas dados necessários, proteger autenticação com Firebase e restringir acesso às informações persistidas."],
        ["RS002", "Uso de credenciais inválidas ou compartilhadas.", "Aplicar validação de login, mensagens de erro controladas e regras de autenticação centralizadas."],
        ["RS003", "Perda de progresso por falha de conexão.", "Tratar indisponibilidade de rede, sincronizar dados quando possível e manter feedback visual ao usuário."],
        ["RS004", "Conteúdo visual ou sonoro gerar desconforto em crianças sensíveis.", "Permitir controle de áudio, utilizar linguagem acolhedora e evitar representações ameaçadoras do atendimento odontológico."],
        ["RS005", "Acesso a telas administrativas ou dados técnicos por usuário infantil.", "Separar responsabilidades de navegação e restringir recursos de manutenção ao ambiente de desenvolvimento ou gestão."],
    ],
)
set_table_font(doc.tables[2], 9)

fill_table(
    doc.tables[3],
    [
        ["Segurança/\nprobabilidade", "Catastrófico\n(1)", "Crítico\n(2)", "Marginal\n(3)", "Insignificante\n(4)"],
        ["Frequente\n(A)", "", "", "RS004", ""],
        ["Provável\n(B)", "", "RS001", "", ""],
        ["Ocasional\n(C)", "", "RS003", "RS002", ""],
        ["Remoto\n(D)", "", "", "RS005", ""],
        ["Improvável\n(E)", "", "", "", ""],
        ["Eliminado\n(F)", "", "", "", ""],
    ],
)
set_table_font(doc.tables[3], 8)

trace_header = ["ID de requisitos"] + [f"RF{i:03d}" for i in range(1, 14)]
trace_rows = [trace_header]
relations = {
    "RF001": {"RF002": "R", "RF009": "D", "RF012": "D"},
    "RF002": {"RF001": "R", "RF010": "D", "RF012": "D"},
    "RF003": {"RF004": "D", "RF005": "D", "RF006": "D", "RF010": "R"},
    "RF004": {"RF003": "D", "RF005": "R", "RF011": "D"},
    "RF005": {"RF003": "D", "RF004": "R", "RF007": "D", "RF011": "D"},
    "RF006": {"RF003": "D", "RF007": "D", "RF011": "D"},
    "RF007": {"RF005": "D", "RF006": "D", "RF009": "R", "RF011": "D"},
    "RF008": {"RF003": "D", "RF011": "D"},
    "RF009": {"RF001": "D", "RF007": "R", "RF002": "D"},
    "RF010": {"RF002": "D", "RF003": "R", "RF004": "D", "RF005": "D", "RF006": "D"},
    "RF011": {"RF004": "D", "RF005": "D", "RF006": "D", "RF007": "D"},
    "RF012": {"RF001": "D", "RF002": "D"},
    "RF013": {"RF003": "D", "RF005": "D", "RF006": "D", "RF010": "D"},
}
for i in range(1, 14):
    rf = f"RF{i:03d}"
    trace_rows.append([rf] + [relations.get(rf, {}).get(f"RF{j:03d}", "") for j in range(1, 14)])
fill_table(doc.tables[4], trace_rows)
set_table_font(doc.tables[4], 6.5)

fill_table(
    doc.tables[5],
    [
        ["REQUISITO FUNCIONAL", "COMPLEXIDADE", "ALI", "AIE", "SE", "EE", "CE", "TOTAL"],
        ["RF001", "MÉDIA", "10", "0", "0", "4", "3", "17"],
        ["RF002", "MÉDIA", "7", "5", "0", "4", "3", "19"],
        ["RF003", "BAIXA", "0", "0", "4", "3", "3", "10"],
        ["RF004", "MÉDIA", "0", "0", "5", "4", "3", "12"],
        ["RF005", "MÉDIA", "0", "0", "5", "4", "3", "12"],
        ["RF006", "MÉDIA", "0", "0", "5", "4", "3", "12"],
        ["RF007", "BAIXA", "7", "0", "4", "3", "3", "17"],
        ["RF008", "BAIXA", "0", "0", "0", "3", "0", "3"],
        ["RF009", "MÉDIA", "10", "5", "4", "4", "3", "26"],
        ["RF010", "BAIXA", "0", "0", "4", "3", "3", "10"],
        ["RF011", "BAIXA", "0", "0", "4", "3", "0", "7"],
        ["RF012", "BAIXA", "7", "0", "0", "3", "0", "10"],
        ["RF013", "BAIXA", "0", "0", "4", "3", "0", "7"],
        ["TOTAL ESTIMADO", "-", "41", "15", "39", "45", "27", "162"],
    ],
)
set_table_font(doc.tables[5], 8)

fill_table(
    doc.tables[6],
    [
        ["Nome: [NOME COMPLETO DO AUTOR 1]\nCPF: [XXX.XXX.XXX-XX]\nQualificação/profissão: [QUALIFICAÇÃO]\nEndereço: [ENDEREÇO COMPLETO]\nCEP: [XXXXX-XXX]\nTelefone: [(XX) XXXXX-XXXX]\nEmail: [email@exemplo.com]"],
        ["Nome: [NOME COMPLETO DO AUTOR 2]\nCPF: [XXX.XXX.XXX-XX]\nQualificação/profissão: [QUALIFICAÇÃO]\nEndereço: [ENDEREÇO COMPLETO]\nCEP: [XXXXX-XXX]\nTelefone: [(XX) XXXXX-XXXX]\nEmail: [email@exemplo.com]"],
        ["Nome: [NOME DO(A) ORIENTADOR(A)]\nCPF: [XXX.XXX.XXX-XX]\nQualificação/profissão: [TITULAÇÃO E FUNÇÃO]\nEndereço: [ENDEREÇO COMPLETO]\nCEP: [XXXXX-XXX]\nTelefone: [(XX) XXXXX-XXXX]\nEmail: [email@exemplo.com]"],
        ["Nome: [NOME DO(A) COORIENTADOR(A), SE HOUVER]\nCPF: [XXX.XXX.XXX-XX]\nQualificação/profissão: [TITULAÇÃO E FUNÇÃO]\nEndereço: [ENDEREÇO COMPLETO]\nCEP: [XXXXX-XXX]\nTelefone: [(XX) XXXXX-XXXX]\nEmail: [email@exemplo.com]"],
        ["Nome: [NOME DO(A) COLABORADOR(A), SE HOUVER]\nCPF: [XXX.XXX.XXX-XX]\nQualificação/profissão: [QUALIFICAÇÃO]\nEndereço: [ENDEREÇO COMPLETO]\nCEP: [XXXXX-XXX]\nTelefone: [(XX) XXXXX-XXXX]\nEmail: [email@exemplo.com]"],
    ],
)

doc.save(OUT)
print(OUT)
